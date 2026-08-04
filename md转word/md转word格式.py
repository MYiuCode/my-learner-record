# -*- coding: utf-8 -*-
"""
Markdown → Word 转换工具
======================
运行后弹出文件选择框，选择 .md 文件后自动在同目录下生成同名 .docx。
转换结果保留原 MD 文档的阅读效果：
  - 标题层级 (H1-H6)
  - 粗体、斜体、删除线、下划线
  - 行内代码、围栏代码块（带语言标签和灰底）
  - 有序/无序/嵌套列表、任务列表
  - 表格（表头加粗 + 灰底）
  - 引用块（左缩进 + 蓝色竖线）
  - 超链接、图片（本地 / base64）
  - 水平线、定义列表、脚注

依赖:
  pip install markdown beautifulsoup4 python-docx

用法:
  python md转word格式.py
"""

import os
import re
import sys
import base64
import tkinter as tk
from io import BytesIO
from tkinter import filedialog, messagebox

# ---------- 第三方库 ----------
import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ============================================================================
#  转换核心
# ============================================================================

class Md2DocxConverter:
    """将 Markdown 文件转换为 Word (.docx) 文件"""

    # ------------------------------------------------------------------ public
    def convert(self, md_path: str, output_path: str | None = None) -> str:
        """转换一个 md 文件为 docx，返回生成的 docx 文件路径。"""
        if output_path is None:
            output_path = self._default_output_path(md_path)

        self._doc = Document()
        self._base_dir = os.path.dirname(os.path.abspath(md_path))
        self._setup_styles()

        html = self._md_to_html(md_path)
        soup = BeautifulSoup(html, "html.parser")
        body = soup.find("body")
        if body:
            self._process_block(list(body.children))
        else:
            self._process_block(list(soup.children))

        self._doc.save(output_path)
        return output_path

    # ----------------------------------------------------------------- helpers
    @staticmethod
    def _default_output_path(md_path: str) -> str:
        base, _ = os.path.splitext(md_path)
        return base + ".docx"

    @staticmethod
    def _read_file(path: str) -> str:
        for enc in ("utf-8", "utf-8-sig", "gbk"):
            try:
                with open(path, "r", encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    def _md_to_html(self, md_path: str) -> str:
        text = self._read_file(md_path)
        extensions = [
            "tables",
            "fenced_code",
            "codehilite",
            "toc",
            "attr_list",
            "def_list",
            "footnotes",
            "md_in_html",
            "sane_lists",
        ]
        html = markdown.markdown(text, extensions=extensions)
        # 后处理：删除线和任务列表（markdown 库默认不支持）
        html = self._postprocess_strikethrough(html)
        html = self._postprocess_tasklist(html)
        return html

    def _postprocess_strikethrough(self, html: str) -> str:
        """将 ~~text~~ 转换为 <del>text</del>，但跳过代码块内的内容。"""
        # 先保护代码块内的内容
        code_blocks = []

        def save_code(match):
            code_blocks.append(match.group(0))
            return f"\x00CODE{len(code_blocks) - 1}\x00"

        # 保护 <pre>、<code> 标签内容
        html = re.sub(r"<pre[^>]*>.*?</pre>", save_code, html, flags=re.DOTALL)
        html = re.sub(r"<code[^>]*>.*?</code>", save_code, html, flags=re.DOTALL)

        # 转换删除线
        html = re.sub(r"~~(.+?)~~", r"<del>\1</del>", html)

        # 还原代码块
        for i, block in enumerate(code_blocks):
            html = html.replace(f"\x00CODE{i}\x00", block)

        return html

    def _postprocess_tasklist(self, html: str) -> str:
        """将列表项中的 [x] / [ ] 转换为复选框。"""
        # 匹配 <li> 标签中的 [x] 或 [ ]
        def replace_checkbox(match):
            prefix = match.group(1)  # <li> 或其他前缀
            checkbox_type = match.group(2)  # x 或空格
            rest = match.group(3)

            checked = ' checked=""' if checkbox_type.strip().lower() == "x" else ""
            checkbox = f'<input type="checkbox"{checked} disabled="">'
            return f"{prefix}{checkbox} {rest}"

        # 匹配 <li...>[x] 或 <li...>[ ] 开头的模式
        html = re.sub(
            r"(<li[^>]*>)\s*\[([ xX])]\s*(.+?)(?=</li>|<ul>|<ol>)",
            replace_checkbox,
            html,
            flags=re.DOTALL,
        )
        return html

    # ============================================================ 样式初始化
    def _setup_styles(self):
        """初始化文档中需要用到的自定义样式"""
        self._ensure_heading_styles()
        self._make_style(
            "CodeBlock", "代码块",
            font_name="Consolas", font_size=Pt(9),
            font_color=RGBColor(0x24, 0x29, 0x2E),
        )
        self._make_style(
            "BlockQuoteText", "引用文字",
            font_size=Pt(10.5),
            font_color=RGBColor(0x6A, 0x73, 0x7D),
        )

    def _ensure_heading_styles(self):
        for level in range(1, 7):
            style_name = f"Heading {level}"
            if style_name not in self._doc.styles:
                self._doc.styles.add_style(style_name, 1)  # 1 = paragraph style

    def _make_style(
        self, style_id: str, style_name: str,
        font_name: str = "微软雅黑",
        font_size: Pt = Pt(10.5),
        font_color: RGBColor | None = None,
    ):
        if style_id in self._doc.styles:
            return
        style = self._doc.styles.add_style(style_id, 1)
        style.font.name = font_name
        style.font.size = font_size
        if font_color:
            style.font.color.rgb = font_color

    # ============================================================ run 工具
    def _add_run(
        self, para, text: str,
        bold=False, italic=False, underline=False, strike=False,
        font_name: str | None = None,
        font_size: Pt | None = None,
        font_color: RGBColor | None = None,
    ):
        """在段落中追加一个 run，并设置格式。"""
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.underline = underline
        if strike:
            rpr = run._element.get_or_add_rPr()
            rpr.append(OxmlElement("w:strike"))
        if font_name:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        if font_size:
            run.font.size = font_size
        if font_color:
            run.font.color.rgb = font_color
        return run

    # ============================================================ 行内处理
    def _add_inline(
        self, para, nodes,
        *, bold=False, italic=False, underline=False, strike=False,
        font_name: str | None = None,
        font_size: Pt | None = None,
        font_color: RGBColor | None = None,
    ):
        """递归处理行内节点（文本 + 行内标签），把内容写入 para。"""
        for node in nodes:
            if isinstance(node, NavigableString):
                text = str(node)
                if text:
                    self._add_run(
                        para, text,
                        bold=bold, italic=italic,
                        underline=underline, strike=strike,
                        font_name=font_name, font_size=font_size,
                        font_color=font_color,
                    )
            elif isinstance(node, Tag):
                tag = node.name
                if tag in ("strong", "b"):
                    self._add_inline(
                        para, node.children,
                        bold=True, italic=italic,
                        underline=underline, strike=strike,
                        font_name=font_name, font_size=font_size,
                        font_color=font_color,
                    )
                elif tag in ("em", "i"):
                    self._add_inline(
                        para, node.children,
                        bold=bold, italic=True,
                        underline=underline, strike=strike,
                        font_name=font_name, font_size=font_size,
                        font_color=font_color,
                    )
                elif tag == "u":
                    self._add_inline(
                        para, node.children,
                        bold=bold, italic=italic,
                        underline=True, strike=strike,
                        font_name=font_name, font_size=font_size,
                        font_color=font_color,
                    )
                elif tag in ("del", "s", "strike"):
                    self._add_inline(
                        para, node.children,
                        bold=bold, italic=italic,
                        underline=underline, strike=True,
                        font_name=font_name, font_size=font_size,
                        font_color=font_color,
                    )
                elif tag == "code":
                    self._add_run(
                        para, node.get_text(),
                        font_name="Consolas",
                        font_size=Pt(9),
                        font_color=RGBColor(0xD7, 0x3A, 0x49),
                    )
                elif tag == "a":
                    href = node.get("href", "")
                    link_text = node.get_text() or href
                    self._add_hyperlink(para, href, link_text)
                elif tag == "img":
                    self._add_image(para, node)
                elif tag == "br":
                    para.add_run().add_break()
                else:
                    # 未知标签，递归处理其子节点
                    self._add_inline(
                        para, node.children,
                        bold=bold, italic=italic,
                        underline=underline, strike=strike,
                        font_name=font_name, font_size=font_size,
                        font_color=font_color,
                    )

    def _add_hyperlink(self, para, url: str, text: str):
        """在段落中插入超链接。"""
        part = para.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")

        color_elem = OxmlElement("w:color")
        color_elem.set(qn("w:val"), "0563C1")
        rpr.append(color_elem)

        u_elem = OxmlElement("w:u")
        u_elem.set(qn("w:val"), "single")
        rpr.append(u_elem)

        sz_elem = OxmlElement("w:sz")
        sz_elem.set(qn("w:val"), "21")  # 10.5pt
        rpr.append(sz_elem)

        new_run.append(rpr)
        text_elem = OxmlElement("w:t")
        text_elem.text = text
        new_run.append(text_elem)
        hyperlink.append(new_run)
        para._p.append(hyperlink)

    def _add_image(self, para, img_tag):
        """在段落中插入图片，找不到文件时显示占位文字。"""
        src = img_tag.get("src", "")
        alt = img_tag.get("alt", "")
        if not src:
            return

        try:
            if src.startswith("data:"):
                match = re.match(r"data:image/(\w+);base64,(.+)", src)
                if match:
                    image_bytes = base64.b64decode(match.group(2))
                    stream = BytesIO(image_bytes)
                    run = para.add_run()
                    run.add_picture(stream, width=Inches(4.5))
                    return
            else:
                img_path = os.path.join(self._base_dir, src)
                if os.path.isfile(img_path):
                    run = para.add_run()
                    run.add_picture(img_path, width=Inches(4.5))
                    return

            # 图片不存在，显示占位
            self._add_run(
                para, f"[图片: {alt or src}]",
                italic=True, font_color=RGBColor(0x99, 0x99, 0x99),
            )
        except Exception:
            self._add_run(
                para, f"[图片加载失败: {src}]",
                italic=True, font_color=RGBColor(0x99, 0x99, 0x99),
            )

    # ============================================================ 块级处理
    def _process_block(self, nodes):
        """递归处理块级节点。"""
        for node in nodes:
            if isinstance(node, NavigableString):
                text = str(node).strip()
                if text:
                    self._doc.add_paragraph(text)
                continue
            if not isinstance(node, Tag):
                continue

            tag = node.name
            if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                self._handle_heading(node)
            elif tag == "p":
                self._handle_p(node)
            elif tag == "pre":
                self._handle_pre(node)
            elif tag == "table":
                self._handle_table(node)
            elif tag == "blockquote":
                self._handle_blockquote(node)
            elif tag in ("ul", "ol"):
                self._handle_list(node, depth=0)
            elif tag == "hr":
                self._handle_hr()
            elif tag in (
                "div", "section", "article", "main",
                "header", "footer", "nav",
            ):
                self._process_block(list(node.children))
            elif tag == "dl":
                self._handle_dl(node)
            # style / script 等标签忽略

    # ----- 标题
    def _handle_heading(self, node):
        level = int(node.name[1])
        para = self._doc.add_heading("", level=level)
        self._add_inline(para, list(node.children))

    # ----- 段落
    def _handle_p(self, node):
        para = self._doc.add_paragraph()
        self._set_paragraph_font(para)
        self._add_inline(para, list(node.children))

    def _set_paragraph_font(self, para):
        """给段落设置默认中文字体。"""
        for run in para.runs:
            if not run.font.name:
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # ----- 代码块
    def _handle_pre(self, node):
        code = node.find("code")
        if code:
            self._handle_code_block(code)
        else:
            self._write_code_block(node.get_text(), lang="")

    def _handle_code_block(self, code_tag):
        classes = code_tag.get("class", [])
        lang = ""
        for cls in classes:
            if cls.startswith("language-"):
                lang = cls[len("language-"):]
                break
        self._write_code_block(code_tag.get_text(), lang=lang)

    def _write_code_block(self, text: str, lang: str = ""):
        """写入一个带灰色底纹的代码块段落。"""
        text = text.rstrip("\n")
        if not text:
            return

        para = self._doc.add_paragraph()
        self._apply_code_block_format(para)

        if lang:
            self._add_run(
                para, f"[{lang}]\n",
                font_name="Consolas", font_size=Pt(7.5),
                font_color=RGBColor(0x99, 0x99, 0x99),
            )

        self._add_run(
            para, text,
            font_name="Consolas", font_size=Pt(9),
            font_color=RGBColor(0x24, 0x29, 0x2E),
        )

    def _apply_code_block_format(self, para):
        """给段落设置代码块样式：灰色背景、等宽字体、合适间距。"""
        pPr = para._p.get_or_add_pPr()

        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F6F8FA")
        shd.set(qn("w:val"), "clear")
        pPr.append(shd)

        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "360")
        ind.set(qn("w:right"), "360")
        pPr.append(ind)

        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "60")
        spacing.set(qn("w:after"), "60")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "exact")
        pPr.append(spacing)

    # ----- 引用块
    def _handle_blockquote(self, node, depth: int = 0):
        indent = 720 + depth * 360
        for child in node.children:
            if isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    para = self._doc.add_paragraph()
                    self._apply_blockquote_format(para, indent)
                    self._add_run(
                        para, text, italic=True,
                        font_color=RGBColor(0x6A, 0x73, 0x7D),
                    )
                continue
            if not isinstance(child, Tag):
                continue

            if child.name == "p":
                para = self._doc.add_paragraph()
                self._apply_blockquote_format(para, indent)
                self._add_inline(
                    para, list(child.children),
                    italic=True, font_color=RGBColor(0x6A, 0x73, 0x7D),
                )
            elif child.name == "blockquote":
                self._handle_blockquote(child, depth + 1)
            elif child.name in ("ul", "ol"):
                self._handle_list(child, depth=0)
            else:
                para = self._doc.add_paragraph()
                self._apply_blockquote_format(para, indent)
                self._add_inline(
                    para, [child],
                    italic=True, font_color=RGBColor(0x6A, 0x73, 0x7D),
                )

    def _apply_blockquote_format(self, para, indent=720):
        """给段落添加引用块格式：左缩进 + 左边灰色竖线。"""
        pPr = para._p.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(indent))
        pPr.append(ind)

        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")    # 宽度 2.25pt
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), "CCCCCC")
        pBdr.append(left)
        pPr.append(pBdr)

    # ----- 列表
    def _handle_list(self, list_node, depth: int = 0):
        """处理 <ul> / <ol>，depth 表示嵌套深度。"""
        ordered = list_node.name == "ol"
        counter = 1
        for child in list_node.children:
            if not isinstance(child, Tag) or child.name != "li":
                continue
            self._handle_list_item(child, ordered=ordered, depth=depth, counter=counter)
            counter += 1

    def _handle_list_item(self, li_node, ordered: bool, depth: int, counter: int):
        para = self._doc.add_paragraph()
        indent_left = 720 + depth * 540
        hanging = 360

        pPr = para._p.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(indent_left))
        ind.set(qn("w:hanging"), str(hanging))
        pPr.append(ind)

        # 任务列表
        checkbox_input = li_node.find("input", {"type": "checkbox"})
        if checkbox_input:
            checked = checkbox_input.has_attr("checked")
            symbol = "☑ " if checked else "☐ "
            self._add_run(para, symbol, font_size=Pt(12))

        # 前缀
        if not checkbox_input:
            if ordered:
                prefix = f"{counter}. "
            else:
                bullets = ["• ", "◦ ", "▪ ", "‣ "]
                prefix = bullets[depth % len(bullets)]
            self._add_run(para, prefix, bold=True, font_size=Pt(10.5))

        # 内容：li 内除了嵌套列表之外的部分
        inline_nodes = []
        nested_list = None
        for child in li_node.children:
            if isinstance(child, Tag) and child.name in ("ul", "ol"):
                nested_list = child
            else:
                inline_nodes.append(child)

        self._add_inline(para, inline_nodes)

        # 嵌套列表
        if nested_list is not None:
            self._handle_list(nested_list, depth=depth + 1)

    # ----- 定义列表
    def _handle_dl(self, node):
        for child in node.children:
            if not isinstance(child, Tag):
                continue
            if child.name == "dt":
                para = self._doc.add_paragraph()
                self._add_inline(para, list(child.children), bold=True)
            elif child.name == "dd":
                para = self._doc.add_paragraph()
                pPr = para._p.get_or_add_pPr()
                ind_elem = OxmlElement("w:ind")
                ind_elem.set(qn("w:left"), "720")
                pPr.append(ind_elem)
                self._add_inline(para, list(child.children))

    # ----- 水平线
    def _handle_hr(self):
        self._doc.add_paragraph()
        para = self._doc.add_paragraph()
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCCC")
        pBdr.append(bottom)
        pPr.append(pBdr)
        self._doc.add_paragraph()

    # ----- 表格
    def _handle_table(self, table):
        all_rows = table.find_all("tr")
        if not all_rows:
            return

        first_row = all_rows[0]
        num_cols = max(1, len(first_row.find_all(["th", "td"])))

        thead = table.find("thead")
        tbody = table.find("tbody")

        if thead:
            header_rows = thead.find_all("tr")
            body_rows = tbody.find_all("tr") if tbody else [
                r for r in all_rows if r not in header_rows
            ]
        else:
            header_rows = []
            body_rows = all_rows

        total_rows = len(header_rows) + len(body_rows)
        if total_rows == 0:
            return

        tbl = self._doc.add_table(rows=total_rows, cols=num_cols)
        tbl.style = "Table Grid"

        row_idx = 0
        for tr in header_rows:
            cells = tr.find_all(["th", "td"])
            for col_idx, cell in enumerate(cells):
                if col_idx >= num_cols:
                    break
                tc = tbl.cell(row_idx, col_idx)
                self._write_table_cell(tc, cell, is_header=True)
            row_idx += 1

        for tr in body_rows:
            cells = tr.find_all(["th", "td"])
            for col_idx, cell in enumerate(cells):
                if col_idx >= num_cols:
                    break
                tc = tbl.cell(row_idx, col_idx)
                self._write_table_cell(tc, cell, is_header=False)
            row_idx += 1

    def _write_table_cell(self, cell, td_tag, is_header: bool):
        """写入表格单元格内容。"""
        for p in cell.paragraphs:
            p.clear()
        para = cell.paragraphs[0]
        if is_header:
            self._apply_header_cell_shading(cell)
            self._add_inline(para, list(td_tag.children), bold=True)
        else:
            self._add_inline(para, list(td_tag.children))

    @staticmethod
    def _apply_header_cell_shading(cell):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F0F0F0")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)


# ============================================================================
#  GUI 入口
# ============================================================================

def main():
    root = tk.Tk()
    root.withdraw()              # 隐藏主窗口
    root.attributes("-topmost", True)

    file_path = filedialog.askopenfilename(
        title="选择 Markdown 文件",
        filetypes=[
            ("Markdown 文件", "*.md *.markdown *.mdown"),
            ("所有文件", "*.*"),
        ],
    )

    if not file_path:
        messagebox.showinfo("提示", "未选择文件，程序已退出。")
        return

    try:
        converter = Md2DocxConverter()
        output_path = converter.convert(file_path)

        messagebox.showinfo(
            "转换成功",
            f"文件已成功转换为 Word 文档！\n\n"
            f"源文件：{os.path.basename(file_path)}\n"
            f"输出文件：{os.path.basename(output_path)}\n"
            f"保存位置：{os.path.dirname(output_path)}",
        )
    except Exception as e:
        messagebox.showerror("转换失败", f"发生错误:\n{e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
