# -*- coding: utf-8 -*-
"""
把 Word 文档里的伪公式标记渲染成真正的上下标
  f_x²  →  fₓ²    （_x 变成下标 x）
  e^y   →  eʸ     （^y 变成上标 y）
  lim_{h→0} → lim 带下标 h→0
  x^{2n} → x 带上标 2n

依赖：pip install python-docx lxml
"""

import os
import re
import sys
import copy
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from threading import Thread

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ═════════════════ 正则：匹配 _^ 上下标标记 ═════════════════
#
# 匹配规则（按优先级）：
#   _{...} 或 ^{...}   → 花括号包裹的内容
#   _(...) 或 ^(...)   → 圆括号包裹的内容
#   _X 或 ^X           → 单个字母/数字/希腊字母
#
# 组1 = _ 或 ^
# 组2 = {...} 内容
# 组3 = (...) 内容
# 组4 = 单字符
MARKER_RE = re.compile(
    r'([_^])'
    r'(?:'
    r'\{([^{}]*)\}'                                    # {...}
    r'|\(([^()]*)\)'                                   # (...)
    r'|([a-zA-Z0-9αβγδεθλμπσφω∂∞′])'                # 单字符
    r')'
)


# ═════════════════ XML 辅助 ═════════════════

def clone_rPr(ref_rPr, vert_align=None):
    """
    复制一份 rPr（文字格式），可选覆盖 vertAlign。
    vert_align: None / 'subscript' / 'superscript'
    """
    new_rPr = OxmlElement('w:rPr')

    if ref_rPr is not None:
        for child in ref_rPr:
            # 跳过原有的 vertAlign，后面统一加
            if child.tag == qn('w:vertAlign'):
                continue
            new_rPr.append(copy.deepcopy(child))

    if vert_align:
        va = OxmlElement('w:vertAlign')
        va.set(qn('w:val'), vert_align)
        new_rPr.append(va)

    return new_rPr


def make_run(text, ref_rPr=None, vert_align=None):
    """创建一个 w:r 元素"""
    r = OxmlElement('w:r')

    if ref_rPr is not None or vert_align:
        r.append(clone_rPr(ref_rPr, vert_align))

    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    return r


def get_ref_rPr(para_elem):
    """取段落中第一个有格式的 run 的 rPr 作为参考"""
    for r in para_elem.findall(qn('w:r')):
        rPr = r.find(qn('w:rPr'))
        if rPr is not None:
            return rPr
    return None


# ═════════════════ 核心：处理段落 ═════════════════

def process_paragraph(para):
    """
    把段落中的 _^ 标记转为真正的上下标。
    返回成功替换的数量。
    """
    text = para.text
    if not text:
        return 0
    if '_' not in text and '^' not in text:
        return 0

    matches = list(MARKER_RE.finditer(text))
    if not matches:
        return 0

    p_elem = para._element
    ref_rPr = get_ref_rPr(p_elem)

    # 备份段落属性
    pPr = p_elem.find(qn('w:pPr'))

    # 删除所有 run
    for r in p_elem.findall(qn('w:r')):
        p_elem.remove(r)

    # 按匹配重建
    last = 0
    count = 0

    for m in matches:
        # 公式标记前的普通文本
        if m.start() > last:
            p_elem.append(make_run(text[last:m.start()], ref_rPr))

        marker = m.group(1)       # _ 或 ^
        content = m.group(2) or m.group(3) or m.group(4)

        if content:
            vert = 'subscript' if marker == '_' else 'superscript'
            p_elem.append(make_run(content, ref_rPr, vert_align=vert))
            count += 1
        else:
            # 异常情况：保留原文
            p_elem.append(make_run(m.group(0), ref_rPr))

        last = m.end()

    # 剩余文本
    if last < len(text):
        p_elem.append(make_run(text[last:], ref_rPr))

    return count


def iter_all_paragraphs(doc):
    """迭代文档中所有段落（正文 + 表格 + 页眉页脚）"""
    yield from doc.paragraphs
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for sec in doc.sections:
        for para in sec.header.paragraphs:
            yield para
        for para in sec.footer.paragraphs:
            yield para


def convert_file(input_path, output_path, log=None):
    """处理整个文档"""
    def _log(msg):
        if log:
            log(msg)

    _log(f"📂 读取：{input_path}")
    doc = Document(input_path)

    total = 0
    para_count = 0

    for para in iter_all_paragraphs(doc):
        n = process_paragraph(para)
        if n > 0:
            total += n
            para_count += 1

    _log(f"📝 处理了 {para_count} 个段落，共 {total} 处上下标")
    _log(f"💾 保存：{output_path}")
    doc.save(output_path)
    return total


# ═════════════════ GUI ═════════════════

class App:
    def __init__(self, root):
        self.root = root
        self.root.title("公式美化：上下标渲染")
        self.root.geometry("620x520")
        self.root.minsize(480, 400)

        self.input_path = tk.StringVar()
        self.output_path = tk.StringVar()
        self._build_ui()

    def _build_ui(self):
        tk.Label(
            self.root, text="公式美化 · 上下标渲染",
            font=("Microsoft YaHei", 15, "bold"),
        ).pack(pady=(14, 4))

        tk.Label(
            self.root,
            text="f_x → fₓ    x^2 → x²    lim_{h→0} → lim 带下标 h→0",
            font=("Microsoft YaHei", 9), fg="#666",
        ).pack(pady=(0, 10))

        # 输入
        f_in = tk.LabelFrame(self.root, text="输入文件", padx=10, pady=8)
        f_in.pack(fill="x", padx=20, pady=5)
        tk.Entry(f_in, textvariable=self.input_path).pack(
            side="left", fill="x", expand=True)
        tk.Button(f_in, text="浏览…", command=self._browse_in, width=8).pack(
            side="right", padx=(6, 0))

        # 输出
        f_out = tk.LabelFrame(self.root, text="输出文件", padx=10, pady=8)
        f_out.pack(fill="x", padx=20, pady=5)
        tk.Entry(f_out, textvariable=self.output_path).pack(
            side="left", fill="x", expand=True)
        tk.Button(f_out, text="浏览…", command=self._browse_out, width=8).pack(
            side="right", padx=(6, 0))

        # 按钮
        self.btn = tk.Button(
            self.root, text="开始转换", command=self._start,
            font=("Microsoft YaHei", 12, "bold"),
            bg="#4CAF50", fg="white", activebackground="#45a049",
            height=2, cursor="hand2",
        )
        self.btn.pack(pady=14, fill="x", padx=40)

        # 日志
        f_log = tk.LabelFrame(self.root, text="日志", padx=8, pady=5)
        f_log.pack(fill="both", expand=True, padx=20, pady=(0, 12))
        self.log_box = scrolledtext.ScrolledText(
            f_log, height=9, font=("Consolas", 9),
            state="disabled", wrap="word",
        )
        self.log_box.pack(fill="both", expand=True)

    def _log(self, msg):
        self.log_box.config(state="normal")
        self.log_box.insert("end", msg + "\n")
        self.log_box.see("end")
        self.log_box.config(state="disabled")
        self.root.update_idletasks()

    def _browse_in(self):
        p = filedialog.askopenfilename(
            title="选择 Word 文档",
            filetypes=[("Word 文档", "*.docx"), ("所有文件", "*.*")],
        )
        if p:
            self.input_path.set(p)
            if not self.output_path.get():
                d, b = os.path.dirname(p), os.path.splitext(os.path.basename(p))[0]
                self.output_path.set(os.path.join(d, f"{b}_美化版.docx"))

    def _browse_out(self):
        p = filedialog.asksaveasfilename(
            title="选择输出路径",
            defaultextension=".docx",
            filetypes=[("Word 文档", "*.docx")],
        )
        if p:
            self.output_path.set(p)

    def _start(self):
        inp = self.input_path.get().strip()
        out = self.output_path.get().strip()

        if not inp or not os.path.isfile(inp):
            messagebox.showwarning("提示", "请选择有效的输入文件")
            return
        if not out:
            messagebox.showwarning("提示", "请选择输出路径")
            return

        self.btn.config(state="disabled", text="转换中…")
        self.log_box.config(state="normal")
        self.log_box.delete("1.0", "end")
        self.log_box.config(state="disabled")

        Thread(target=self._run, args=(inp, out), daemon=True).start()

    def _run(self, inp, out):
        try:
            n = convert_file(inp, out, log=self._log)
            self._log(f"\n✅ 完成！共渲染 {n} 处上下标")
            self.root.after(0, lambda: messagebox.showinfo(
                "完成", f"转换完成！\n共渲染 {n} 处上下标"))
        except Exception as e:
            error_msg = str(e)
            self._log(f"\n❌ 失败：{error_msg}")
            import traceback
            self._log(traceback.format_exc())
            self.root.after(0, lambda msg=error_msg: messagebox.showerror("错误", msg))
        finally:
            self.btn.config(state="normal", text="开始转换")


def main():
    root = tk.Tk()
    App(root)
    root.mainloop()


if __name__ == "__main__":
    main()
